from datetime import datetime
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_word_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(fld_end)


def add_page_numbers(doc: Document, skip_first: bool = True) -> None:
    from docx.oxml import OxmlElement
    for i, section in enumerate(doc.sections):
        if skip_first and i == 0:
            continue
        footer = section.footer
        # Add page number into its own paragraph to avoid clearing other footer content
        p = footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def add_header(doc: Document, header_text: str, skip_first: bool = True) -> None:
    for i, section in enumerate(doc.sections):
        if skip_first and i == 0:
            continue
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = header_text
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.italic = True


def add_footer_date(doc: Document, date_text: str, skip_first: bool = True) -> None:
    for i, section in enumerate(doc.sections):
        if skip_first and i == 0:
            continue
        footer = section.footer
        p = footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        r = p.add_run(date_text)
        r.font.size = Pt(10)
        r.italic = True


def format_quote_paragraph(p) -> None:
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.italic = True
        r.font.size = Pt(10)


def _setup_styles(doc: Document, base_font_size_pt: int = 12) -> None:
    # Normal paragraphs: justify
    try:
        normal = doc.styles["Normal"]
        normal.font.size = Pt(base_font_size_pt)
        normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    except Exception:
        pass

    # Headings: center, bold, +2pt relative to body
    for style_name in ["Heading 1", "Heading 2"]:
        try:
            st = doc.styles[style_name]
            st.font.bold = True
            st.font.size = Pt(base_font_size_pt + 2)
            st.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception:
            continue

    # Title page heading
    try:
        title_style = doc.styles["Title"]
        title_style.font.bold = True
        title_style.font.size = Pt(base_font_size_pt + 4)
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        pass


BANNED_PATTERNS = [
    re.compile(r"\bлгбт\+?\b", re.IGNORECASE),
    re.compile(r"\blgbt\w*\b", re.IGNORECASE),
]


def _contains_banned(text: str | None) -> bool:
    if not text:
        return False
    return any(p.search(text) for p in BANNED_PATTERNS)


def build_document(country: str, title: str, sections: list, distribution_table: list | None = None) -> Document:
    doc = Document()
    _setup_styles(doc, base_font_size_pt=12)
    created_date = datetime.now().strftime('%d.%m.%Y')

    # Title page
    title_para = doc.add_heading(f"Проблемное поле жителей {translate_country_ru(country)}: тематический анализ на основании материалов социальных сетей и электронных медиа", 0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para = doc.add_paragraph(f"Дата формирования отчёта: {created_date}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for r in date_para.runs:
        r.font.size = Pt(14)

    # New page for TOC and content
    doc.add_section(0)  # new page

    # TOC right after title page
    doc.add_heading("Оглавление", level=1)
    add_word_toc(doc)
    doc.add_page_break()

    # Research description
    doc.add_heading("Описание исследования", level=1)
    desc = (
        f"Настоящий отчёт представляет результаты кабинетного разведочного анализа публичных сообщений жителей {translate_country_ru(country)} "
        "в социальных сетях и электронных медиа. Цель — выявить ключевые проблемы, волнующие рядовых граждан, на основе фактов и конкретных событий.\n\n"
        f"Дата создания отчёта: {created_date}\n"
        "Методология: качественный контент‑анализ с использованием LLM для структурирования информации; минимизация анализа в пользу фактов; редактор‑агент удаляет повторы, проверяет логику и цитаты.\n"
        "Источники: сообщения и материалы, собранные из открытых источников (соцсети, СМИ)."
    )
    p_desc = doc.add_paragraph(desc)
    p_desc.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Distribution table
    if distribution_table:
        doc.add_heading("Распределение проблем", level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Проблема"
        hdr[1].text = "Количество"
        for row in distribution_table:
            r = table.add_row().cells
            r[0].text = str(row[0])
            r[1].text = str(row[1])

    # Sections
    for section in sections:
        title_text = section.get("title")
        if _contains_banned(title_text):
            continue

        # Filter blocks for banned content and empty sources
        blocks_in = section.get("blocks", [])
        filtered_blocks = []
        for block in blocks_in:
            btype = block.get("type")
            if btype in {"paragraph", "quote", "subheading"}:
                text_val = block.get("text")
                # Do not render a 'Источники' subheading here; sources are handled via 'sources' block
                if btype == "subheading" and text_val and text_val.strip().lower().startswith("источники"):
                    continue
                if _contains_banned(text_val):
                    continue
                filtered_blocks.append(block)
            elif btype == "sources":
                items = [str(u) for u in block.get("items", []) if not _contains_banned(str(u))]
                if items:
                    filtered_blocks.append({"type": "sources", "items": items})

        if not filtered_blocks:
            # Skip empty sections after filtering
            continue

        doc.add_heading(title_text, level=1)  # Heading 1
        for block in filtered_blocks:
            if block["type"] == "subheading":
                doc.add_heading(block["text"], level=2)  # Heading 2
            elif block["type"] == "paragraph":
                p = doc.add_paragraph(block["text"])  # Normal
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif block["type"] == "quote":
                qp = doc.add_paragraph(block["text"])  # Normal then format
                format_quote_paragraph(qp)
            elif block["type"] == "sources":
                items = block.get("items", [])
                if items:
                    doc.add_heading("Источники", level=2)
                    for url in items:
                        p = doc.add_paragraph(str(url))
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_header(doc, "Проблемное поле жителей {0}: тематический анализ на основании материалов социальных сетей и электронных медиа".format(translate_country_ru(country)), skip_first=True)
    add_footer_date(doc, f"Дата создания отчёта: {created_date}", skip_first=True)
    add_page_numbers(doc, skip_first=True)
    return doc


def translate_country_ru(country: str) -> str:
    mapping = {
        "Tanzania": "Танзании",
    }
    return mapping.get(country, country)

